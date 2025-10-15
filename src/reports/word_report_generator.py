#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Генератор отчетов в формате Word (.docx)
"""

from datetime import datetime
from typing import Dict, List, Optional
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

class WordReportGenerator:
    """Генератор отчетов в формате Word"""
    
    def __init__(self):
        self.document = None
        self.title_style = None
        self.heading_style = None
        self.normal_style = None
    
    def create_report(self, program_state_manager, db_manager, output_path: str = None) -> str:
        """Создает отчет в формате Word"""
        
        # Создаем документ
        self.document = Document()
        self._setup_styles()
        
        # Генерируем отчет
        self._add_title()
        self._add_summary_section(program_state_manager, db_manager)
        self._add_networks_section(program_state_manager, db_manager)
        self._add_action_history_section(program_state_manager)
        self._add_metrics_section(program_state_manager)
        
        # Сохраняем файл
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"ics_report_{timestamp}.docx"
        
        # Убеждаемся, что директория существует
        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
        
        self.document.save(output_path)
        return output_path
    
    def _setup_styles(self):
        """Настраивает стили документа"""
        # Стиль заголовка
        self.title_style = self.document.styles.add_style('CustomTitle', 1)  # 1 = WD_STYLE_TYPE.PARAGRAPH
        self.title_style.font.name = 'Arial'
        self.title_style.font.size = Pt(18)
        self.title_style.font.bold = True
        self.title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.title_style.paragraph_format.space_after = Pt(12)
        
        # Стиль подзаголовков
        self.heading_style = self.document.styles.add_style('CustomHeading', 1)  # 1 = WD_STYLE_TYPE.PARAGRAPH
        self.heading_style.font.name = 'Arial'
        self.heading_style.font.size = Pt(14)
        self.heading_style.font.bold = True
        self.heading_style.paragraph_format.space_before = Pt(12)
        self.heading_style.paragraph_format.space_after = Pt(6)
        
        # Стиль обычного текста
        self.normal_style = self.document.styles.add_style('CustomNormal', 1)  # 1 = WD_STYLE_TYPE.PARAGRAPH
        self.normal_style.font.name = 'Arial'
        self.normal_style.font.size = Pt(11)
        self.normal_style.paragraph_format.space_after = Pt(6)
    
    def _add_title(self):
        """Добавляет заголовок отчета"""
        title = self.document.add_paragraph()
        title.style = self.title_style
        title_run = title.add_run("ОТЧЕТ ИКС АНАЛИЗАТОРА СИСТЕМЫ")
        
        subtitle = self.document.add_paragraph()
        subtitle.style = self.normal_style
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run(f"Дата и время генерации: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}")
        
        # Добавляем пустую строку
        self.document.add_paragraph()
    
    def _add_summary_section(self, program_state_manager, db_manager):
        """Добавляет раздел с общей информацией"""
        heading = self.document.add_paragraph()
        heading.style = self.heading_style
        heading.add_run("1. ОБЩАЯ ИНФОРМАЦИЯ")
        
        # Получаем информацию о состоянии
        status_info = program_state_manager.get_status_info()
        networks_status = program_state_manager.get_networks_status(db_manager)
        
        # Создаем таблицу с общей информацией
        table = self.document.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Заголовки таблицы
        table.cell(0, 0).text = "Параметр"
        table.cell(0, 1).text = "Значение"
        
        # Заполняем таблицу
        table.cell(1, 0).text = "Текущее состояние программы"
        table.cell(1, 1).text = status_info['state_display']
        
        table.cell(2, 0).text = "Время выполнения"
        table.cell(2, 1).text = status_info['runtime_display']
        
        table.cell(3, 0).text = "Количество пауз"
        table.cell(3, 1).text = str(status_info['pause_count'])
        
        table.cell(4, 0).text = "Общее время на паузе"
        table.cell(4, 1).text = status_info['total_pause_time_display']
        
        table.cell(5, 0).text = "Последнее действие"
        if status_info['last_action_time']:
            table.cell(5, 1).text = status_info['last_action_time'].strftime('%d.%m.%Y %H:%M:%S')
        else:
            table.cell(5, 1).text = "Нет данных"
        
        # Добавляем пустую строку
        self.document.add_paragraph()
    
    def _add_networks_section(self, program_state_manager, db_manager):
        """Добавляет раздел с информацией о сетях"""
        heading = self.document.add_paragraph()
        heading.style = self.heading_style
        heading.add_run("2. ИНФОРМАЦИЯ О СЕТЯХ")
        
        # Получаем статус сетей
        networks_status = program_state_manager.get_networks_status(db_manager)
        
        if 'error' in networks_status:
            error_para = self.document.add_paragraph()
            error_para.style = self.normal_style
            error_para.add_run(f"Ошибка получения информации о сетях: {networks_status['error']}")
            return
        
        # Общая статистика
        stats_para = self.document.add_paragraph()
        stats_para.style = self.normal_style
        stats_para.add_run(f"Всего сетей в базе данных: {networks_status['total_networks']}")
        
        if networks_status['total_networks'] > 0:
            # Проверяем тип networks
            networks_list = networks_status.get('networks', [])
            
            if not isinstance(networks_list, list):
                networks_list = []
            
            # Создаем таблицу с сетями
            table = self.document.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Заголовки таблицы
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'ID'
            hdr_cells[1].text = 'Название'
            hdr_cells[2].text = 'Дата создания'
            hdr_cells[3].text = 'Статус'
            
            # Заполняем таблицу
            for network in networks_list:
                if isinstance(network, dict):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(network.get('id', 'N/A'))
                    network_name = network.get('name')
                    row_cells[1].text = network_name if network_name is not None else 'N/A'
                    created_at = network.get('created_at')
                    row_cells[2].text = created_at if created_at is not None else 'N/A'
                    status = network.get('status')
                    row_cells[3].text = status if status is not None else 'N/A'
                else:
                    pass  # Пропускаем не-словарь
        else:
            no_networks_para = self.document.add_paragraph()
            no_networks_para.style = self.normal_style
            no_networks_para.add_run("Сети не найдены")
        
        # Текущая сеть
        status_info = program_state_manager.get_status_info()
        current_network_name = status_info.get('current_network_name')
        if current_network_name is not None:
            current_para = self.document.add_paragraph()
            current_para.style = self.normal_style
            current_network_id = status_info.get('current_network_id', 'N/A')
            current_para.add_run(f"Текущая активная сеть: {current_network_name} (ID: {current_network_id})")
        
        # Добавляем пустую строку
        self.document.add_paragraph()
    
    def _add_action_history_section(self, program_state_manager):
        """Добавляет раздел с историей действий"""
        heading = self.document.add_paragraph()
        heading.style = self.heading_style
        heading.add_run("3. ИСТОРИЯ ДЕЙСТВИЙ")
        
        # Получаем журнал действий
        action_log = program_state_manager.get_action_log(limit=100)
        
        if action_log and isinstance(action_log, list):
            # Создаем таблицу с историей
            table = self.document.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Заголовки таблицы
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Время'
            hdr_cells[1].text = 'Действие'
            hdr_cells[2].text = 'Детали'
            hdr_cells[3].text = 'Сеть'
            
            # Заполняем таблицу
            for log_entry in reversed(action_log):  # Показываем последние действия первыми
                if isinstance(log_entry, dict):
                    row_cells = table.add_row().cells
                    row_cells[0].text = log_entry.get('timestamp', 'N/A')
                    row_cells[1].text = log_entry.get('action', 'N/A')
                    row_cells[2].text = log_entry.get('details', 'N/A')
                    network_name = log_entry.get('network_name')
                    row_cells[3].text = network_name if network_name is not None else '-'
                else:
                    pass  # Пропускаем не-словарь в action_log
        else:
            no_actions_para = self.document.add_paragraph()
            no_actions_para.style = self.normal_style
            no_actions_para.add_run("История действий пуста")
        
        # Добавляем пустую строку
        self.document.add_paragraph()
    
    def _add_metrics_section(self, program_state_manager):
        """Добавляет раздел с метриками"""
        heading = self.document.add_paragraph()
        heading.style = self.heading_style
        heading.add_run("4. МЕТРИКИ ПРОГРАММЫ")
        
        # Получаем метрики
        status_info = program_state_manager.get_status_info()
        metrics = status_info['metrics']
        
        # Создаем таблицу с метриками
        table = self.document.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Заголовки таблицы
        table.cell(0, 0).text = "Метрика"
        table.cell(0, 1).text = "Значение"
        
        # Заполняем таблицу
        table.cell(1, 0).text = "Всего создано сетей"
        table.cell(1, 1).text = str(metrics['total_networks_created'])
        
        table.cell(2, 0).text = "Всего удалено сетей"
        table.cell(2, 1).text = str(metrics['total_networks_deleted'])
        
        table.cell(3, 0).text = "Всего запущено симуляций"
        table.cell(3, 1).text = str(metrics['total_simulations_run'])
        
        table.cell(4, 0).text = "Общее время работы"
        table.cell(4, 1).text = self._format_duration(metrics['total_runtime_seconds'])
        
        # Добавляем пустую строку
        self.document.add_paragraph()
        
        # Добавляем заключение
        conclusion = self.document.add_paragraph()
        conclusion.style = self.heading_style
        conclusion.add_run("ЗАКЛЮЧЕНИЕ")
        
        conclusion_text = self.document.add_paragraph()
        conclusion_text.style = self.normal_style
        
        conclusion_content = f"""
Данный отчет содержит полную информацию о работе ИКС Анализатора Системы.
Программа находилась в состоянии: {status_info['state_display']}.
Общее время работы составило: {status_info['runtime_display']}.
Было создано {metrics['total_networks_created']} сетей и выполнено {metrics['total_simulations_run']} симуляций.
"""
        
        conclusion_text.add_run(conclusion_content.strip())
    
    def _format_duration(self, seconds: float) -> str:
        """Форматирует длительность в читаемый вид"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        
        if hours > 0:
            return f"{hours:02d}:{minutes:02d}:{secs:02d}"
        else:
            return f"{minutes:02d}:{secs:02d}"
